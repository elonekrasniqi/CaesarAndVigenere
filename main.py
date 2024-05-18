import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog
from functools import partial
import os
import docx
import fitz  # PyMuPDF

def caesar_encrypt(input_text, shift):
    encrypted_text = ""
    for char in input_text:
        if char.isalpha():
            if char.isupper():
                encrypted_text += chr((ord(char) + shift - 65) % 26 + 65)
            else:
                encrypted_text += chr((ord(char) + shift - 97) % 26 + 97)
        else:
            encrypted_text += char
    return encrypted_text

def caesar_decrypt(encrypted_text, shift):
    decrypted_text = ""
    for char in encrypted_text:
        if char.isalpha():
            if char.isupper():
                decrypted_text += chr((ord(char) - shift - 65) % 26 + 65)
            else:
                decrypted_text += chr((ord(char) - shift - 97) % 26 + 97)
        else:
            decrypted_text += char
    return decrypted_text

def vigenere_encrypt(text, key):
    result = ""
    key_index = 0
    for char in text:
        if char.isalpha():
            shift = ord(key[key_index].upper()) - 65
            if char.isupper():
                result += chr((ord(char) - 65 + shift) % 26 + 65)
            else:
                result += chr((ord(char) - 97 + shift) % 26 + 97)
            key_index = (key_index + 1) % len(key)
        else:
            result += char
    return result


def vigenere_decrypt(text, key):
    result = ""
    key_index = 0
    for char in text:
        if char.isalpha():
            shift = ord(key[key_index].upper()) - 65
            if char.isupper():
                result += chr((ord(char) - 65 - shift + 26) % 26 + 65)
            else:
                result += chr((ord(char) - 97 - shift + 26) % 26 + 97)
            key_index = (key_index + 1) % len(key)
        else:
            result += char
    return result



def read_pdf(file_path):
    text = ""
    doc = fitz.open(file_path)
    for page_num in range(doc.page_count):
        page = doc.load_page(page_num)
        text += page.get_text()
    return text


def save_pdf(file_path, text):
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), text)
    doc.save(file_path)


def process_file(algorithm, operation, file_path, shift_or_key):
    input_text = ""
    output_text = ""
    if file_path.endswith('.txt'):
        with open(file_path, "r") as file:
            input_text = file.read()
    elif file_path.endswith('.docx'):
        doc = docx.Document(file_path)
        input_text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
    elif file_path.endswith('.pdf'):
        input_text = read_pdf(file_path)

    if algorithm == "Caesar":
        if operation == "encrypt":
            output_text = caesar_encrypt(input_text, shift_or_key)
        else:
            output_text = caesar_decrypt(input_text, shift_or_key)
    elif algorithm == "Vigenere":
        if not shift_or_key.isalpha():
            messagebox.showerror("Error", "The key must be alphabetic.")
            return None, None
        if operation == "encrypt":
            output_text = vigenere_encrypt(input_text, shift_or_key)
        else:
            output_text = vigenere_decrypt(input_text, shift_or_key)





    output_file_name = os.path.splitext(os.path.basename(file_path))[0] + f"_{algorithm.lower()}_{operation}.txt"
    if file_path.endswith('.txt'):
        output_file_name = os.path.splitext(os.path.basename(file_path))[0] + f"_{algorithm.lower()}_{operation}.txt"
    elif file_path.endswith('.docx'):
        output_file_name = os.path.splitext(os.path.basename(file_path))[0] + f"_{algorithm.lower()}_{operation}.docx"
    elif file_path.endswith('.pdf'):
        output_file_name = os.path.splitext(os.path.basename(file_path))[0] + f"_{algorithm.lower()}_{operation}.pdf"

    return output_text, output_file_name



def perform_operation(algorithm, operation, from_file, text_box, temp_file_path=None):
    global save_path_selected
    save_path_selected = False  # Reset the flag each time this function is called

    if from_file:
        file_path = filedialog.askopenfilename(filetypes=[("Text files", "*.txt"), ("Word files", "*.docx"), ("PDF files", "*.pdf")])
        if not file_path:
            messagebox.showerror("Error", "No file selected.")
            return
    else:
        file_path = temp_file_path

    while not save_path_selected:
        save_path = filedialog.askdirectory(title="Select Save Location")
        if not save_path:
            if messagebox.askretrycancel("Error", "No save location selected. Do you want to retry?"):
                continue
            else:
                return
        else:
            save_path_selected = True




    if algorithm == "Caesar":
        shift = simpledialog.askinteger("Shift Value", "Enter the shift value (must be a number):")
        if shift is None:
            messagebox.showerror("Error", "No shift value entered.")
            return
        output_text, output_file_name = process_file(algorithm, operation, file_path, shift)

    elif algorithm == "Vigenere":
        key = simpledialog.askstring("Key", "Enter the encryption key (must be alphabetic):")
        if key is None:
            messagebox.showerror("Error", "No key entered.")
            return
        if not key.isalpha():
            messagebox.showerror("Error", "The key must be alphabetic.")
            return
        output_text, output_file_name = process_file(algorithm, operation, file_path, key)

    else:
        messagebox.showerror("Error", "Invalid algorithm.")
        return

    output_file_path = os.path.join(save_path, output_file_name)
    
    if output_file_path.endswith('.docx'):
        doc = docx.Document()
        doc.add_paragraph(output_text)
        doc.save(output_file_path)
    elif output_file_path.endswith('.pdf'):
        save_pdf(output_file_path, output_text)
    else:
        with open(output_file_path, "w") as file:
            file.write(output_text)

    messagebox.showinfo("Success", f"File {operation}ed successfully.")
    if not from_file and temp_file_path:
        os.remove(temp_file_path)




def show_operation_selection(algorithm):
    global selected_algorithm
    selected_algorithm = algorithm
    algorithm_frame.pack_forget()
    operation_frame.pack(padx=10, pady=5)


def show_file_selection(operation):
    global selected_operation
    selected_operation = operation
    operation_frame.pack_forget()
    file_frame.pack(padx=10, pady=5)


def handle_file_choice(has_file):
    global selected_has_file
    selected_has_file = has_file
    file_frame.pack_forget()
    if has_file == "no":
        text_frame.pack(padx=10, pady=5)
        process_button.pack(padx=10, pady=5)
    else:
        perform_operation(selected_algorithm, selected_operation, True, text_box)



def save_text_to_file():
    global temp_file_path
    text = text_box.get("1.0", tk.END).strip()
    if not text:
        messagebox.showerror("Error", "No text entered.")
        return
    temp_file_path = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("Text files", "*.txt"), ("Word files", "*.docx"), ("PDF files", "*.pdf")])
    if not temp_file_path:
        messagebox.showerror("Error", "No save location selected.")
        return
    if temp_file_path.endswith('.docx'):
        doc = docx.Document()
        doc.add_paragraph(text)
        doc.save(temp_file_path)
    elif temp_file_path.endswith('.pdf'):
        save_pdf(temp_file_path, text)
    else:
        with open(temp_file_path, "w") as file:
            file.write(text)
    messagebox.showinfo("Success", f"Text saved to {temp_file_path}")
    perform_operation(selected_algorithm, selected_operation, False, text_box, temp_file_path)



def start_process():
    perform_operation(selected_algorithm, selected_operation, selected_has_file == "yes", text_box)



def main():
    global text_box, process_button, temp_file_path, save_path_selected
    global algorithm_frame, operation_frame, file_frame, text_frame
    global selected_algorithm, selected_operation, selected_has_file

    root = tk.Tk()
    root.title("Cipher")

    main_frame = tk.Frame(root)
    main_frame.pack(padx=10, pady=10)

    algorithm_frame = tk.Frame(main_frame)
    algorithm_frame.pack(padx=10, pady=5)
    algorithm_label = tk.Label(algorithm_frame, text="Select an Algorithm:", font=("Arial", 12, "bold"))
    algorithm_label.pack()
    caesar_button = tk.Button(algorithm_frame, text="Caesar", command=lambda: show_operation_selection("Caesar"), bg="blue", fg="white", font=("Arial", 12, "bold"))
    caesar_button.pack(side=tk.LEFT, padx=5)
    vigenere_button = tk.Button(algorithm_frame, text="Vigenere", command=lambda: show_operation_selection("Vigenere"), bg="green", fg="white", font=("Arial", 12, "bold"))
    vigenere_button.pack(side=tk.LEFT, padx=5)

    operation_frame = tk.Frame(main_frame)
    operation_label = tk.Label(operation_frame, text="Select Operation:", font=("Arial", 12, "bold"))
    operation_label.pack()
    encrypt_button = tk.Button(operation_frame, text="Encrypt", command=lambda: show_file_selection("encrypt"), bg="blue", fg="white", font=("Arial", 12, "bold"))
    encrypt_button.pack(side=tk.LEFT, padx=5)
    decrypt_button = tk.Button(operation_frame, text="Decrypt", command=lambda: show_file_selection("decrypt"), bg="green", fg="white", font=("Arial", 12, "bold"))
    decrypt_button.pack(side=tk.LEFT, padx=5)

    file_frame = tk.Frame(main_frame)
    file_label = tk.Label(file_frame, text="Do you have a file to process?", font=("Arial", 12, "bold"))
    file_label.pack()
    file_yes_button = tk.Button(file_frame, text="Yes", command=lambda: handle_file_choice("yes"), bg="blue", fg="white", font=("Arial", 12, "bold"))
    file_yes_button.pack(side=tk.LEFT, padx=5)
    file_no_button = tk.Button(file_frame, text="No", command=lambda: handle_file_choice("no"), bg="green", fg="white", font=("Arial", 12, "bold"))
    file_no_button.pack(side=tk.LEFT, padx=5)

    text_frame = tk.Frame(main_frame)
    text_box_label = tk.Label(text_frame, text="Enter your text:", font=("Arial", 12, "bold"))
    text_box_label.pack()
    text_box = tk.Text(text_frame, height=10, width=50)
    text_box.pack(pady=5)
    save_text_button = tk.Button(text_frame, text="Save Text to File", command=save_text_to_file, bg="orange", fg="white", font=("Arial", 12, "bold"))
    save_text_button.pack(pady=5)

    process_button = tk.Button(main_frame, text="Process", command=start_process, bg="blue", fg="white", font=("Arial", 12, "bold"))

    root.mainloop()

if __name__ == "__main__":
    main()
